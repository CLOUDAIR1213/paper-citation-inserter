from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document

from citation_utils import find_citation_markers, is_probable_heading, is_reference_heading, normalize_reference_key
from extract_bibliography import extract_bibliography_entries


def sentence_for_marker(text: str, start: int, end: int) -> str:
    left_candidates = [text.rfind(mark, 0, start) for mark in "。！？.!?"]
    right_candidates = [text.find(mark, end) for mark in "。！？.!?"]
    left = max(left_candidates) + 1
    right_positions = [pos for pos in right_candidates if pos != -1]
    right = min(right_positions) + 1 if right_positions else len(text)
    return text[left:right].strip()


def load_reference_notes(path: Path | None) -> list[dict]:
    if path is None:
        return []
    data = json.loads(path.read_text(encoding="utf-8"))
    return data.get("entries", [])


def match_note(reference_text: str, notes: list[dict]) -> dict | None:
    reference_key = normalize_reference_key(reference_text)
    reference_lower = reference_text.lower()
    for note in notes:
        title = str(note.get("title_guess") or "").strip().lower()
        raw = str(note.get("raw_text") or "").strip().lower()
        if title and title in reference_lower:
            return note
        if title and title in reference_key:
            return note
        if reference_key and reference_key[:80] in raw:
            return note
    return None


def previous_text(paragraphs, index: int) -> str:
    for paragraph in reversed(paragraphs[:index]):
        text = paragraph.text.strip()
        if text and not is_reference_heading(text):
            return text
    return ""


def next_text(paragraphs, index: int, bibliography_start: int | None) -> str:
    limit = bibliography_start if bibliography_start is not None else len(paragraphs)
    for paragraph in paragraphs[index + 1 : limit]:
        text = paragraph.text.strip()
        if text and not is_reference_heading(text):
            return text
    return ""


def extract_citation_contexts(docx_path: Path, notes_path: Path | None = None) -> dict:
    doc = Document(str(docx_path))
    entries, bibliography_start = extract_bibliography_entries(docx_path)
    reference_map = {entry.original_number: entry.text for entry in entries if entry.original_number is not None}
    notes = load_reference_notes(notes_path)
    contexts = []
    current_section = ""

    scan_end = bibliography_start if bibliography_start is not None else len(doc.paragraphs)
    for index, paragraph in enumerate(doc.paragraphs[:scan_end]):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if is_probable_heading(style_name, text):
            current_section = text
            continue
        if not text:
            continue

        for marker in find_citation_markers(text):
            mapped_references = []
            for number in marker["numbers"]:
                reference_text = reference_map.get(number)
                mapped_references.append(
                    {
                        "number": number,
                        "missing": reference_text is None,
                        "reference": reference_text,
                        "matched_note": match_note(reference_text, notes) if reference_text else None,
                    }
                )

            contexts.append(
                {
                    "paragraph_index": index,
                    "section": current_section,
                    "marker": marker["raw"],
                    "citation_numbers": marker["numbers"],
                    "sentence": sentence_for_marker(text, marker["start"], marker["end"]),
                    "paragraph": text,
                    "previous_paragraph": previous_text(doc.paragraphs, index),
                    "next_paragraph": next_text(doc.paragraphs, index, bibliography_start),
                    "mapped_references": mapped_references,
                }
            )

    return {
        "source": str(docx_path),
        "reference_notes": str(notes_path) if notes_path else None,
        "bibliography_start_paragraph": bibliography_start,
        "contexts": contexts,
    }


def write_markdown(data: dict, path: Path) -> None:
    lines = [
        "# Citation Contexts",
        "",
        f"- Source: `{data['source']}`",
        f"- Reference notes: `{data['reference_notes']}`" if data.get("reference_notes") else "- Reference notes: None",
        f"- Citation markers: {len(data['contexts'])}",
        "",
    ]
    for item in data["contexts"]:
        lines.extend(
            [
                f"## Paragraph {item['paragraph_index']} {item['marker']}",
                "",
                f"- Section: {item['section'] or 'N/A'}",
                f"- Sentence: {item['sentence']}",
                f"- Previous: {item['previous_paragraph'] or 'N/A'}",
                f"- Next: {item['next_paragraph'] or 'N/A'}",
                "",
                "| Number | Status | Reference | Note |",
                "| ---: | --- | --- | --- |",
            ]
        )
        for reference in item["mapped_references"]:
            note = reference.get("matched_note") or {}
            note_label = note.get("source_id") or note.get("title_guess") or ""
            status = "missing reference" if reference["missing"] else "matched"
            ref_text = (reference.get("reference") or "").replace("|", "\\|")
            lines.append(f"| {reference['number']} | {status} | {ref_text} | {note_label} |")
        lines.append("")

    path.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract citation contexts for semantic alignment review.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--reference-notes", "-n", type=Path)
    parser.add_argument("--output", "-o", type=Path, default=Path("citation_contexts.json"))
    parser.add_argument("--markdown", "-m", type=Path, default=Path("citation_contexts.md"))
    args = parser.parse_args()

    data = extract_citation_contexts(args.docx, args.reference_notes)
    args.output.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown(data, args.markdown)
    print(f"Wrote {args.output}")
    print(f"Wrote {args.markdown}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
