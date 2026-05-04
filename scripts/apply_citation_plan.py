from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path

from docx import Document

from citation_utils import normalize_reference_key, strip_leading_reference_number
from extract_bibliography import extract_bibliography_entries
from repair_citation_order import repair_document


def bibliography_start_index(doc: Document) -> int | None:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() in {"参考文献", "references", "bibliography"}:
            return index
    return None


def append_reference_entries(doc: Document, entries_to_add: list[str]) -> list[int]:
    start = bibliography_start_index(doc)
    if start is None:
        doc.add_paragraph("参考文献")

    existing, _ = extract_bibliography_entries_from_doc(doc)
    next_number = max([entry["number"] for entry in existing if entry["number"] is not None] or [0]) + 1
    assigned_numbers = []
    for entry_text in entries_to_add:
        doc.add_paragraph(f"[{next_number}] {strip_leading_reference_number(entry_text)}")
        assigned_numbers.append(next_number)
        next_number += 1
    return assigned_numbers


def extract_bibliography_entries_from_doc(doc: Document) -> tuple[list[dict], int | None]:
    from citation_utils import is_reference_heading, parse_reference_line

    entries = []
    start = None
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if start is None:
            if is_reference_heading(text):
                start = index
            continue
        if not text:
            continue
        number, body = parse_reference_line(text)
        entries.append({"number": number, "text": body if number is not None else text})
    return entries, start


def insert_marker(paragraph_text: str, sentence: str | None, marker: str) -> str:
    if sentence and sentence in paragraph_text:
        return paragraph_text.replace(sentence, f"{sentence}{marker}", 1)
    return paragraph_text.rstrip() + marker


def apply_plan_to_copy(source_path: Path, plan_path: Path, working_path: Path) -> dict:
    shutil.copyfile(source_path, working_path)
    plan = json.loads(plan_path.read_text(encoding="utf-8"))
    insertions = [item for item in plan.get("insertions", []) if item.get("confidence", "high") == "high"]
    candidates = plan.get("candidates", [])

    doc = Document(str(working_path))
    existing_entries, _ = extract_bibliography_entries(working_path)
    known_references = {normalize_reference_key(entry.text): entry.original_number for entry in existing_entries}

    references_to_add: list[str] = []
    insertion_refs: list[str] = []
    for item in insertions:
        reference_text = item.get("reference_text", "").strip()
        if not reference_text:
            continue
        key = normalize_reference_key(reference_text)
        insertion_refs.append(reference_text)
        if key not in known_references and key not in {normalize_reference_key(ref) for ref in references_to_add}:
            references_to_add.append(reference_text)

    new_numbers = append_reference_entries(doc, references_to_add)
    added_number_by_key = {
        normalize_reference_key(ref): number for ref, number in zip(references_to_add, new_numbers, strict=False)
    }

    for item in insertions:
        paragraph_index = item.get("paragraph_index")
        reference_text = item.get("reference_text", "").strip()
        if paragraph_index is None or not reference_text:
            continue
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            continue
        key = normalize_reference_key(reference_text)
        old_number = known_references.get(key) or added_number_by_key.get(key)
        if old_number is None:
            continue
        paragraph = doc.paragraphs[paragraph_index]
        paragraph.text = insert_marker(paragraph.text, item.get("sentence"), f"[{old_number}]")

    doc.save(str(working_path))
    return {"insertions": insertions, "candidates": candidates}


def append_plan_notes(report_path: Path, plan_notes: dict) -> None:
    lines = report_path.read_text(encoding="utf-8").splitlines()
    lines.extend(["", "## Semantic Insertions", ""])
    if plan_notes["insertions"]:
        for item in plan_notes["insertions"]:
            lines.append(
                f"- Paragraph {item.get('paragraph_index')}: {item.get('reference_text', '').strip()} "
                f"(reason: {item.get('reason', 'not provided')})"
            )
    else:
        lines.append("- None")

    lines.extend(["", "## 待确认候选", ""])
    if plan_notes["candidates"]:
        for item in plan_notes["candidates"]:
            lines.append(
                f"- Paragraph {item.get('paragraph_index')}: {item.get('reference_text', '').strip()} "
                f"(reason: {item.get('reason', 'not provided')})"
            )
    else:
        lines.append("- None")
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Apply an explicit citation insertion plan, then repair final citation order.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("plan", type=Path)
    parser.add_argument("--output", "-o", type=Path)
    parser.add_argument("--report", "-r", type=Path)
    parser.add_argument("--work", type=Path)
    args = parser.parse_args()

    output = args.output or args.docx.with_name(f"{args.docx.stem}_with_citations.docx")
    report = args.report or output.with_name(f"{output.stem}_report.md")
    work = args.work or output.with_name(f"{output.stem}_before_repair.docx")

    try:
        notes = apply_plan_to_copy(args.docx, args.plan, work)
        repair_document(work, output, report)
        append_plan_notes(report, notes)
    finally:
        if work.exists() and args.work is None:
            work.unlink()

    print(f"Wrote {output}")
    print(f"Wrote {report}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
