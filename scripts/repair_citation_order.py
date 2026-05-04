from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from citation_utils import (
    CITATION_PATTERN,
    ReferenceEntry,
    find_citation_markers,
    format_citation_numbers,
    is_reference_heading,
)
from extract_bibliography import extract_bibliography_entries


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def bibliography_start_index(doc: Document) -> int | None:
    for index, paragraph in enumerate(doc.paragraphs):
        if is_reference_heading(paragraph.text):
            return index
    return None


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = text


def build_reference_map(entries: list[ReferenceEntry]) -> dict[int, ReferenceEntry]:
    return {entry.original_number: entry for entry in entries if entry.original_number is not None}


def rebuild_text_citations(
    text: str,
    reference_map: dict[int, ReferenceEntry],
    old_to_new: dict[int, int],
    usage_log: list[dict],
    missing_numbers: set[int],
    paragraph_index: int,
) -> str:
    def replace(match: re.Match) -> str:
        old_numbers = find_citation_markers(match.group(0))[0]["numbers"]
        rebuilt_numbers: list[int] = []
        missing_in_marker: list[int] = []

        for old_number in old_numbers:
            if old_number not in reference_map:
                missing_numbers.add(old_number)
                missing_in_marker.append(old_number)
                continue
            if old_number not in old_to_new:
                old_to_new[old_number] = len(old_to_new) + 1
            rebuilt_numbers.append(old_to_new[old_number])
            usage_log.append(
                {
                    "paragraph_index": paragraph_index,
                    "old_number": old_number,
                    "new_number": old_to_new[old_number],
                    "reference": reference_map[old_number].text,
                    "original_marker": match.group(0),
                }
            )

        if rebuilt_numbers and not missing_in_marker:
            return format_citation_numbers(rebuilt_numbers)
        if rebuilt_numbers and missing_in_marker:
            return format_citation_numbers(rebuilt_numbers) + "[missing:" + ",".join(str(n) for n in missing_in_marker) + "]"
        return match.group(0)

    return CITATION_PATTERN.sub(replace, text)


def rewrite_bibliography(doc: Document, reference_map: dict[int, ReferenceEntry], old_to_new: dict[int, int]) -> None:
    start = bibliography_start_index(doc)
    if start is None:
        doc.add_paragraph("参考文献")
        start = len(doc.paragraphs) - 1

    for paragraph in list(doc.paragraphs[start + 1 :]):
        delete_paragraph(paragraph)

    ordered = sorted(old_to_new.items(), key=lambda item: item[1])
    for old_number, new_number in ordered:
        doc.add_paragraph(reference_map[old_number].renumbered(new_number))


def write_report(
    report_path: Path,
    source_path: Path,
    output_path: Path,
    usage_log: list[dict],
    reference_map: dict[int, ReferenceEntry],
    old_to_new: dict[int, int],
    missing_numbers: set[int],
) -> None:
    used_old_numbers = set(old_to_new)
    unused = [entry for old_number, entry in sorted(reference_map.items()) if old_number not in used_old_numbers]

    lines = [
        "# Citation Repair Report",
        "",
        f"- Source: `{source_path}`",
        f"- Output: `{output_path}`",
        f"- Unique cited references: {len(old_to_new)}",
        f"- Unused bibliography references: {len(unused)}",
        f"- Missing bibliography entries cited in text: {len(missing_numbers)}",
        "",
        "## Renumbered Citations",
        "",
        "| Paragraph | Old | New | Reference |",
        "| --- | ---: | ---: | --- |",
    ]

    seen_rows = set()
    for item in usage_log:
        key = (item["paragraph_index"], item["old_number"], item["new_number"])
        if key in seen_rows:
            continue
        seen_rows.add(key)
        reference = item["reference"].replace("|", "\\|")
        lines.append(f"| {item['paragraph_index']} | {item['old_number']} | {item['new_number']} | {reference} |")

    lines.extend(["", "## 文末存在但正文未引用", ""])
    if unused:
        for entry in unused:
            lines.append(f"- {entry.text}")
    else:
        lines.append("- None")

    lines.extend(["", "## 正文引用但文末缺失", ""])
    if missing_numbers:
        for number in sorted(missing_numbers):
            lines.append(f"- 原正文引用编号 [{number}] 未在文末参考文献中找到对应条目，需人工补全。")
    else:
        lines.append("- None")

    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def repair_document(source_path: Path, output_path: Path, report_path: Path) -> dict:
    entries, _ = extract_bibliography_entries(source_path)
    reference_map = build_reference_map(entries)
    doc = Document(str(source_path))
    start = bibliography_start_index(doc)
    scan_end = start if start is not None else len(doc.paragraphs)
    old_to_new: dict[int, int] = {}
    usage_log: list[dict] = []
    missing_numbers: set[int] = set()

    for index, paragraph in enumerate(doc.paragraphs[:scan_end]):
        text = paragraph.text
        if not text:
            continue
        rebuilt = rebuild_text_citations(text, reference_map, old_to_new, usage_log, missing_numbers, index)
        if rebuilt != text:
            replace_paragraph_text(paragraph, rebuilt)

    rewrite_bibliography(doc, reference_map, old_to_new)
    doc.save(str(output_path))
    write_report(report_path, source_path, output_path, usage_log, reference_map, old_to_new, missing_numbers)
    return {
        "source": str(source_path),
        "output": str(output_path),
        "report": str(report_path),
        "renumbered": old_to_new,
        "missing": sorted(missing_numbers),
    }


def default_output_path(source: Path) -> Path:
    return source.with_name(f"{source.stem}_citation_repaired.docx")


def main() -> int:
    parser = argparse.ArgumentParser(description="Repair citation order and bibliography order in a .docx paper.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output", "-o", type=Path)
    parser.add_argument("--report", "-r", type=Path)
    args = parser.parse_args()

    output = args.output or default_output_path(args.docx)
    report = args.report or output.with_name(f"{output.stem}_report.md")
    result = repair_document(args.docx, output, report)
    print(f"Wrote {result['output']}")
    print(f"Wrote {result['report']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
