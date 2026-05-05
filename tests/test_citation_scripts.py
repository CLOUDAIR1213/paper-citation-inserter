from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from repair_citation_order import repair_document
from apply_citation_plan import apply_plan_to_copy
from extract_citation_contexts import extract_citation_contexts


def paragraph_texts(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(str(path)).paragraphs]


def citation_runs(path: Path) -> list[tuple[str, bool | None]]:
    runs = []
    for paragraph in Document(str(path)).paragraphs:
        for run in paragraph.runs:
            if "[" in run.text and "]" in run.text:
                runs.append((run.text, run.font.superscript))
    return runs


class CitationScriptTests(unittest.TestCase):
    def make_docx(self, path: Path) -> None:
        doc = Document()
        doc.add_heading("绪论", level=1)
        doc.add_paragraph("第二篇文献先在正文出现[2]，后面又引用第一篇文献[1]。")
        doc.add_paragraph("第二篇文献再次出现时仍应复用编号[2]。")
        doc.add_paragraph("合并引用需要展开检查[1,2]，范围引用也要展开[1-3]。")
        doc.add_paragraph("这一句引用了不存在的文献[9]。")
        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] Zhang A. First source[J]. Journal, 2020, 1(1): 1-5.")
        doc.add_paragraph("[2] Li B. Second source[J]. Journal, 2021, 2(1): 6-9.")
        doc.add_paragraph("[3] Wang C. Range source[J]. Journal, 2022, 3(1): 10-12.")
        doc.add_paragraph("[4] Zhao D. Unused source[J]. Journal, 2023, 4(1): 13-15.")
        doc.save(str(path))

    def test_repair_order_and_report(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "paper.docx"
            output = Path(tmp) / "paper_repaired.docx"
            report = Path(tmp) / "report.md"
            self.make_docx(source)

            repair_document(source, output, report)

            texts = paragraph_texts(output)
            self.assertIn("第二篇文献先在正文出现[1]，后面又引用第一篇文献[2]。", texts)
            self.assertIn("第二篇文献再次出现时仍应复用编号[1]。", texts)
            self.assertIn("合并引用需要展开检查[1-2]，范围引用也要展开[1-3]。", texts)
            self.assertIn("[1] Li B. Second source", "\n".join(texts))
            self.assertIn("[2] Zhang A. First source", "\n".join(texts))
            self.assertNotIn("[4] Zhao D. Unused source", "\n".join(texts))
            self.assertTrue(any(text == "[1]" and superscript is True for text, superscript in citation_runs(output)))
            self.assertTrue(any(text == "[2]" and superscript is True for text, superscript in citation_runs(output)))
            self.assertTrue(any(text == "[1-2]" and superscript is True for text, superscript in citation_runs(output)))
            self.assertTrue(any(text == "[1-3]" and superscript is True for text, superscript in citation_runs(output)))

            report_text = report.read_text(encoding="utf-8")
            self.assertIn("Zhao D. Unused source", report_text)
            self.assertIn("原正文引用编号 [9]", report_text)

    def test_apply_plan_adds_reference_then_repairs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "paper.docx"
            working = Path(tmp) / "working.docx"
            plan = Path(tmp) / "plan.json"
            self.make_docx(source)
            plan.write_text(
                json.dumps(
                    {
                        "insertions": [
                            {
                                "paragraph_index": 1,
                                "reference_text": "Chen D. New source[J]. Journal, 2024, 4(1): 13-18.",
                                "reason": "supports the opening claim",
                                "confidence": "high",
                            }
                        ],
                        "candidates": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            apply_plan_to_copy(source, plan, working)

            texts = paragraph_texts(working)
            self.assertTrue(any("Chen D. New source" in text for text in texts))
            self.assertTrue(any("[5]" in text for text in texts))
            self.assertTrue(any(text == "[5]" and superscript is True for text, superscript in citation_runs(working)))

    def test_extract_citation_contexts_maps_references_and_missing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "paper.docx"
            notes = Path(tmp) / "reference_entries.json"
            self.make_docx(source)
            notes.write_text(
                json.dumps(
                    {
                        "entries": [
                            {
                                "source_id": "R1",
                                "title_guess": "Second source",
                                "raw_text": "Li B. Second source[J]. Journal, 2021, 2(1): 6-9.",
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            data = extract_citation_contexts(source, notes)

            markers = [item["marker"] for item in data["contexts"]]
            self.assertIn("[2]", markers)
            self.assertIn("[1,2]", markers)
            self.assertIn("[1-3]", markers)
            first = next(item for item in data["contexts"] if item["marker"] == "[2]")
            self.assertEqual(first["mapped_references"][0]["reference"].split()[1], "Li")
            self.assertEqual(first["mapped_references"][0]["matched_note"]["source_id"], "R1")

            missing = next(item for item in data["contexts"] if item["marker"] == "[9]")
            self.assertTrue(missing["mapped_references"][0]["missing"])


if __name__ == "__main__":
    unittest.main()
